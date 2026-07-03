import json
import time
from io import BytesIO
from pathlib import Path

from docx import Document
from ebooklib import epub
import pytest

import echodraft_api.ingestion as ingestion
from echodraft_api.cleaning import CleaningPipeline
from echodraft_api.ingestion import IngestionError, IngestionService


def project_id(client) -> str:
    response = client.post("/api/v1/projects", json={"title": "Import Test", "rightsStatus": "declared"})
    assert response.status_code == 201
    return response.json()["id"]


def wait_for_job(client, job_id: str) -> dict:
    for _ in range(50):
        job = client.get(f"/api/v1/jobs/{job_id}").json()
        if job["status"] in {"succeeded", "failed"}:
            return job
        time.sleep(0.02)
    raise AssertionError("Import job did not finish")


def docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("A DOCX manuscript line.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def epub_bytes(tmp_path: Path) -> bytes:
    book = epub.EpubBook()
    book.set_identifier("sample")
    book.set_title("Sample")
    book.set_language("en")
    chapter = epub.EpubHtml(title="One", file_name="one.xhtml", lang="en")
    chapter.content = "<h1>One</h1><p>An EPUB manuscript line.</p>"
    book.add_item(chapter)
    book.toc = (chapter,)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    target = tmp_path / "sample.epub"
    epub.write_epub(str(target), book)
    return target.read_bytes()


def pdf_bytes(text: str) -> bytes:
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, payload in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{number} 0 obj\n".encode() + payload + b"\nendobj\n")
    xref = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    result.extend(b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:]))
    result.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    return bytes(result)


def import_bytes(client, project: str, name: str, data: bytes, mime: str = "application/octet-stream") -> dict:
    response = client.post(
        f"/api/v1/projects/{project}/source/import",
        files={"file": (name, data, mime)},
        data={"rightsAcknowledged": "true"},
    )
    assert response.status_code == 202, response.text
    return wait_for_job(client, response.json()["id"])


def test_txt_import_creates_deterministic_canonical_artifacts_and_manifest(client) -> None:
    project = project_id(client)
    job = import_bytes(client, project, "story.txt", b"A\r\n\r\nA\r\n\r\nAn   ending.\r\n")
    assert job["status"] == "succeeded"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    assert source["preview"] == "A\n\nAn ending.\n"
    manifest = json.loads(Path(source["manifestPath"]).read_text())
    assert manifest["payload"]["canonicalChecksum"]
    assert any(item["message"].startswith("Duplicated") for item in source["warnings"])


def test_cleaning_pipeline_preserves_numeric_chapter_markers() -> None:
    result = CleaningPipeline().clean(
        "1\n\nChapter One\n\nPage 12\n\nA wrap-\naround repair and a wrapped\nsentence."
    )

    assert result.text == "1\n\nChapter One\n\nA wraparound repair and a wrapped sentence."
    assert "line_page_marker" in {change.change_type for change in result.changes}
    assert "hyphenation_repair" in {change.change_type for change in result.changes}
    assert "line_wrap_merge" in {change.change_type for change in result.changes}


def test_ingestion_removes_page_markers_from_canonical_text(client) -> None:
    project = project_id(client)
    text = (
        "Ooh! I felt a wiggle that time.\n\n"
        "<!-- Page 9 -->\n\n"
        "Open! My eyelids creep up."
    )
    job = import_bytes(client, project, "page-marker.md", text.encode())
    assert job["status"] == "succeeded"

    source = client.get(f"/api/v1/projects/{project}/source").json()
    canonical = Path(source["canonicalPath"]).read_text(encoding="utf-8")
    assert source["preview"] == "Ooh! I felt a wiggle that time.\n\nOpen! My eyelids creep up.\n"
    assert "<!-- Page 9 -->" not in canonical
    assert source["warnings"][0]["message"].startswith("Canonical cleaning applied")
    source_by_id = client.get(f"/api/v1/sources/{source['id']}").json()
    assert source_by_id["preview"] == source["preview"]
    runs = client.get(f"/api/v1/sources/{source['id']}/cleaning-runs").json()
    assert runs[0]["status"] == "succeeded"
    assert Path(runs[0]["manifestPath"]).is_file()
    issues = client.get(f"/api/v1/sources/{source['id']}/cleaning-issues").json()
    assert issues[0]["issueType"] == "html_page_marker"
    assert issues[0]["status"] == "applied"


def test_cleaning_issues_flag_suspicious_tokens_and_can_be_resolved(client) -> None:
    project = project_id(client)
    job = import_bytes(client, project, "ocr-ish.txt", b"Chapter 1\n\nThe rn0on shimmered.")
    assert job["status"] == "succeeded"

    source = client.get(f"/api/v1/projects/{project}/source").json()
    issues = client.get(f"/api/v1/sources/{source['id']}/cleaning-issues").json()
    suspicious = next(item for item in issues if item["issueType"] == "suspicious_ocr_token")
    assert suspicious["status"] == "open"
    assert suspicious["severity"] == "warning"

    response = client.patch(
        f"/api/v1/cleaning-issues/{suspicious['id']}",
        json={"status": "resolved", "resolvedByUser": True},
    )
    assert response.status_code == 200
    assert response.json()["status"] == "resolved"
    assert response.json()["resolvedByUser"] is True


def test_markdown_docx_and_epub_import(client, tmp_path: Path) -> None:
    for name, data in [("story.md", b"# Heading\n\nText"), ("story.docx", docx_bytes()), ("story.epub", epub_bytes(tmp_path))]:
        project = project_id(client)
        job = import_bytes(client, project, name, data)
        assert job["status"] == "succeeded", job["errorMessage"]
        assert client.get(f"/api/v1/projects/{project}/source").json()["preview"]


def test_pdf_import_extracts_text_and_preserves_manifest(client, monkeypatch) -> None:
    monkeypatch.setattr(ingestion, "resolve_system_tool", lambda _: None)
    project = project_id(client)
    job = import_bytes(
        client,
        project,
        "story.pdf",
        pdf_bytes("A PDF manuscript line with enough readable characters."),
        "application/pdf",
    )
    assert job["status"] == "succeeded", job["errorMessage"]
    source = client.get(f"/api/v1/projects/{project}/source").json()
    assert source["mimeType"] == "application/pdf"
    assert "A PDF manuscript line with enough readable characters." in source["preview"]
    assert Path(source["originalPath"]).suffix == ".pdf"


def test_pdf_import_creates_page_review_records(client, app, monkeypatch) -> None:
    monkeypatch.setattr(ingestion, "resolve_system_tool", lambda _: None)
    project = project_id(client)
    job = import_bytes(
        client,
        project,
        "story.pdf",
        pdf_bytes("A page-aware PDF manuscript line with enough readable characters."),
        "application/pdf",
    )
    assert job["status"] == "succeeded", job["errorMessage"]
    source = client.get(f"/api/v1/projects/{project}/source").json()

    pages = client.get(f"/api/v1/sources/{source['id']}/pages")
    assert pages.status_code == 200
    payload = pages.json()
    assert len(payload) == 1
    assert payload[0]["pageNumber"] == 1
    assert payload[0]["extractionMethod"] == "embedded_text"
    assert "page-aware PDF manuscript" in payload[0]["preview"]
    assert Path(payload[0]["embeddedTextPath"]).is_file()
    assert Path(payload[0]["selectedTextPath"]).is_file()
    spans = app.state.container.source_artifacts.spans(source["id"])
    assert len(spans) == 1
    assert spans[0].page_number == 1


def test_pdf_ocr_candidates_and_failures(app, monkeypatch, tmp_path: Path) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class FakeReader:
        is_encrypted = False

        def __init__(self, pages: list[FakePage]) -> None:
            self.pages = pages

    service = IngestionService(app.state.container)
    monkeypatch.setattr(
        ingestion,
        "PdfReader",
        lambda _: FakeReader([FakePage("A readable page with sufficient extracted text."), FakePage("")]),
    )
    monkeypatch.setattr(IngestionService, "_require_ocr_tools", staticmethod(lambda: None))
    monkeypatch.setattr(
        IngestionService, "_ocr_page", staticmethod(lambda _pdf, _root, page: f"OCR page {page} text.")
    )
    text, warnings = service._extract_pdf(tmp_path / "mixed.pdf")
    assert "A readable page with sufficient extracted text." in text and "OCR page 2 text." in text
    assert text.split("\n\n") == [
        "A readable page with sufficient extracted text.",
        "OCR page 2 text.",
    ]
    assert warnings[0].source_range == "page 2"

    monkeypatch.undo()
    monkeypatch.setattr(ingestion, "PdfReader", lambda _: FakeReader([FakePage("")]))
    monkeypatch.setattr(ingestion, "resolve_system_tool", lambda _: None)
    with pytest.raises(IngestionError, match="pdftoppm"):
        service._extract_pdf(tmp_path / "needs-ocr.pdf")

    monkeypatch.undo()
    monkeypatch.setattr(ingestion, "PdfReader", lambda _: FakeReader([FakePage("")]))
    monkeypatch.setattr(
        ingestion,
        "resolve_system_tool",
        lambda command: "/usr/bin/pdftoppm" if command == "pdftoppm" else None,
    )
    with pytest.raises(IngestionError, match="Tesseract"):
        service._extract_pdf(tmp_path / "needs-tesseract.pdf")

    monkeypatch.undo()
    monkeypatch.setattr(
        ingestion, "PdfReader", lambda _: FakeReader([FakePage("") for _ in range(151)])
    )
    with pytest.raises(IngestionError, match="150 pages"):
        service._extract_pdf(tmp_path / "too-many.pdf")


def test_scanned_pdf_path_uses_mocked_ocr_for_every_page(app, monkeypatch, tmp_path: Path) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakeReader:
        is_encrypted = False
        pages = [FakePage(), FakePage(), FakePage()]

    service = IngestionService(app.state.container)
    monkeypatch.setattr(ingestion, "PdfReader", lambda _: FakeReader())
    monkeypatch.setattr(IngestionService, "_require_ocr_tools", staticmethod(lambda: None))
    monkeypatch.setattr(
        IngestionService,
        "_ocr_page",
        staticmethod(lambda _pdf, _root, page: f"Recognized scanned page {page}."),
    )

    text, warnings = service._extract_pdf(tmp_path / "scanned.pdf")

    assert text == (
        "Recognized scanned page 1.\n\n"
        "Recognized scanned page 2.\n\n"
        "Recognized scanned page 3."
    )
    assert [warning.source_range for warning in warnings] == ["page 1", "page 2", "page 3"]


def test_pdf_v2_scanned_page_records_ocr_artifacts(
    client, app, monkeypatch: pytest.MonkeyPatch
) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakeReader:
        is_encrypted = False
        pages = [FakePage()]

    def fake_render(_pdf_path: Path, output_root: Path, page_number: int) -> Path:
        output_root.mkdir(parents=True, exist_ok=True)
        image_path = output_root / f"page_{page_number:04d}.png"
        image_path.write_bytes(b"png")
        return image_path

    def fake_ocr(
        _image_path: Path, output_root: Path, page_number: int
    ) -> tuple[str, Path, Path, float]:
        output_root.mkdir(parents=True, exist_ok=True)
        text_path = output_root / f"page_{page_number:04d}.txt"
        json_path = output_root / f"page_{page_number:04d}.json"
        text_path.write_text("Recognized v2 scanned text.", encoding="utf-8")
        json_path.write_text("{}", encoding="utf-8")
        return "Recognized v2 scanned text.", text_path, json_path, 0.82

    project = project_id(client)
    service = IngestionService(app.state.container)
    source_id = service.stage(
        project,
        "scan.pdf",
        "application/pdf",
        b"%PDF-1.4 fake",
        ingestion.PARSER_VERSION,
    )
    source = app.state.container.sources.get(source_id)
    assert source
    monkeypatch.setattr(ingestion, "PdfReader", lambda _: FakeReader())
    monkeypatch.setattr(
        ingestion,
        "resolve_system_tool",
        lambda command: f"/usr/bin/{command}" if command in {"pdftoppm", "tesseract"} else None,
    )
    monkeypatch.setattr(IngestionService, "_render_pdf_page", staticmethod(fake_render))
    monkeypatch.setattr(IngestionService, "_ocr_page_image", staticmethod(fake_ocr))

    text, warnings = service._extract_pdf_v2(Path(source.original_path), source_id, project)

    assert text == "Recognized v2 scanned text."
    assert warnings[0].source_range == "page 1"
    pages = app.state.container.source_artifacts.pages(source_id)
    assert len(pages) == 1
    assert pages[0].extraction_method == "ocr"
    runs = app.state.container.source_artifacts.ocr_runs(source_id)
    assert len(runs) == 1 and runs[0].status == "succeeded"
    results = app.state.container.source_artifacts.ocr_results(runs[0].id)
    assert len(results) == 1
    assert Path(results[0].text_path).is_file()

def test_ocr_warning_and_failed_parse_preserve_original(client) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "ocr.txt", "Bad � text".encode())["status"] == "succeeded"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    assert source["warnings"][0]["severity"] == "warning"
    project = project_id(client)
    failed = import_bytes(client, project, "broken.docx", b"not a docx")
    assert failed["status"] == "failed"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    assert source["status"] == "failed"
    assert Path(source["originalPath"]).exists()
    project = project_id(client)
    failed = import_bytes(client, project, "broken.pdf", b"not a pdf", "application/pdf")
    assert failed["status"] == "failed"


def docx_heading_bytes() -> bytes:
    document = Document()
    document.add_heading("The Arrival", level=1)
    document.add_paragraph("Body of the first chapter.")
    document.add_heading("A Quiet Scene", level=2)
    document.add_paragraph("Body of the scene.")
    document.add_heading("The Departure", level=1)
    document.add_paragraph("Body of the second chapter.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def multi_item_epub_bytes(tmp_path: Path) -> bytes:
    book = epub.EpubBook()
    book.set_identifier("multi")
    book.set_title("Multi")
    book.set_language("en")
    first = epub.EpubHtml(title="The Arrival", file_name="c1.xhtml", lang="en", uid="arrival")
    first.content = "<h1>The Arrival</h1><p>Body one.</p>"
    second = epub.EpubHtml(title="The Departure", file_name="c2.xhtml", lang="en", uid="departure")
    second.content = "<h1>The Departure</h1><p>Body two.</p>"
    third = epub.EpubHtml(title="Untitled", file_name="c3.xhtml", lang="en", uid="return")
    third.content = "<h1>The Return</h1><p>Body three.</p>"
    book.add_item(first)
    book.add_item(second)
    book.add_item(third)
    book.toc = (first, second)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", first, second, third]
    target = tmp_path / "multi.epub"
    epub.write_epub(str(target), book)
    return target.read_bytes()


def _signals_from_manifest(source: dict) -> list[dict]:
    manifest = json.loads(Path(source["manifestPath"]).read_text())
    signals_path = manifest["payload"]["structureSignalsPath"]
    assert signals_path, "manifest payload is missing structureSignalsPath"
    assert f"/sources/{source['id']}/structure_signals/" in signals_path.replace("\\", "/")
    return json.loads(Path(signals_path).read_text())


def test_docx_heading_styles_emit_ordered_chapter_signals(client) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "headings.docx", docx_heading_bytes())["status"] == "succeeded"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    signals = _signals_from_manifest(source)
    assert [(item["title"], item["sourceKind"], item["level"]) for item in signals] == [
        ("The Arrival", "docx_heading", 1),
        ("A Quiet Scene", "docx_heading", 2),
        ("The Departure", "docx_heading", 1),
    ]
    assert signals[0]["anchorText"] == "The Arrival"


def test_epub_spine_order_and_toc_emit_chapter_signals(client, tmp_path: Path) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "multi.epub", multi_item_epub_bytes(tmp_path))["status"] == "succeeded"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    signals = _signals_from_manifest(source)
    # TOC entries first (0.95), then the first-h1 of a spine item not covered by TOC (0.8).
    assert [(item["title"], item["sourceKind"]) for item in signals] == [
        ("The Arrival", "epub_toc"),
        ("The Departure", "epub_toc"),
        ("The Return", "epub_spine"),
    ]
    assert signals[-1]["confidence"] == 0.8


def test_txt_import_records_no_structure_signals(client) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "plain.txt", b"A plain paragraph.\n")["status"] == "succeeded"
    source = client.get(f"/api/v1/projects/{project}/source").json()
    manifest = json.loads(Path(source["manifestPath"]).read_text())
    assert manifest["payload"]["structureSignalsPath"] is None


def test_reparse_docx_regenerates_signals_for_new_source(client) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "headings.docx", docx_heading_bytes())["status"] == "succeeded"
    first = client.get(f"/api/v1/projects/{project}/source").json()
    first_signals = json.loads(
        Path(json.loads(Path(first["manifestPath"]).read_text())["payload"]["structureSignalsPath"]).read_text()
    )
    response = client.post(
        f"/api/v1/projects/{project}/source/reparse", json={"parserVersion": "ingestion-0.1.1"}
    )
    assert response.status_code == 202
    assert wait_for_job(client, response.json()["id"])["status"] == "succeeded"
    second = client.get(f"/api/v1/projects/{project}/source").json()
    assert second["id"] != first["id"]
    second_path = json.loads(Path(second["manifestPath"]).read_text())["payload"]["structureSignalsPath"]
    assert f"/sources/{second['id']}/" in second_path.replace("\\", "/")
    second_signals = json.loads(Path(second_path).read_text())
    assert [item["title"] for item in second_signals] == [item["title"] for item in first_signals]

    extract_job = client.post(
        f"/api/v1/projects/{project}/structure/extract", json={"maxSegmentChars": 200}
    ).json()
    assert wait_for_job(client, extract_job["id"])["status"] == "succeeded"
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters] == ["The Arrival", "The Departure"]


def test_reparse_preserves_prior_manifest(client) -> None:
    project = project_id(client)
    assert import_bytes(client, project, "story.txt", b"One paragraph.")["status"] == "succeeded"
    first = client.get(f"/api/v1/projects/{project}/source").json()
    response = client.post(f"/api/v1/projects/{project}/source/reparse", json={"parserVersion": "ingestion-0.1.1"})
    assert response.status_code == 202
    assert wait_for_job(client, response.json()["id"])["status"] == "succeeded"
    second = client.get(f"/api/v1/projects/{project}/source").json()
    assert Path(first["manifestPath"]).exists()
    assert first["manifestPath"] != second["manifestPath"]
