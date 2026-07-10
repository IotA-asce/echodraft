import json
import time
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import echodraft_api.cast_discovery as cast_discovery_module
import echodraft_api.structure as structure_module
import echodraft_api.structure_v2 as structure_v2
import pytest
from docx import Document
from ebooklib import epub
from fastapi.testclient import TestClient
from echodraft_api.main import create_app
from echodraft_api.structure_parsing import (
    ChapterSignal,
    StructureCompiler,
    TextAtom,
    validate_atom_offsets,
)


def wait_for_job(client, job_id: str) -> dict:
    for _ in range(60):
        job = client.get(f"/api/v1/jobs/{job_id}").json()
        if job["status"] in {"succeeded", "failed"}:
            return job
        time.sleep(0.02)
    raise AssertionError("job did not finish")


def project_with_source(client, text: str) -> str:
    project = client.post("/api/v1/projects", json={"title": "Structure", "rightsStatus": "declared"}).json()["id"]
    job = client.post(f"/api/v1/projects/{project}/source/import", files={"file": ("book.txt", text.encode(), "text/plain")}, data={"rightsAcknowledged": "true"}).json()
    assert wait_for_job(client, job["id"])["status"] == "succeeded"
    return project


def extract(client, project: str) -> None:
    job = client.post(f"/api/v1/projects/{project}/structure/extract", json={"maxSegmentChars": 120}).json()
    assert wait_for_job(client, job["id"])["status"] == "succeeded"


def atom_ids_from_prompt(prompt: str) -> list[str]:
    return [line.split()[1] for line in prompt.splitlines() if line.startswith("ATOM ")]


def all_segments(client, project: str) -> list[dict]:
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    scenes = [
        scene
        for chapter in chapters
        for scene in client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()
    ]
    return [
        segment
        for scene in scenes
        for segment in client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    ]


def warning_codes(client, project: str) -> set[str]:
    warnings = client.get(f"/api/v1/projects/{project}/structure-warnings").json()
    return {str(warning["evidence"].get("code") or "") for warning in warnings}


def structure_warnings(client, project: str) -> list[dict]:
    return client.get(f"/api/v1/projects/{project}/structure-warnings").json()


def issue_codes(client, project: str) -> set[str]:
    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    return {str(issue["metadata"].get("code") or "") for issue in issues}


def cast_graph_mentions(client, project: str):
    return client.app.state.container.cast_graph.mentions(project)


def cast_graph_decisions(client, project: str):
    return client.app.state.container.cast_graph.decisions(project)


def casting_manifest(client, project: str) -> dict:
    record = client.app.state.container.projects.get(project)
    assert record is not None
    path = Path(record.artifact_path) / "manifests" / "casting_manifest.json"
    return json.loads(path.read_text(encoding="utf-8"))


def test_heading_scene_and_sentence_safe_segments(client) -> None:
    project = project_with_source(client, "Chapter 1: Arrival\n\nMara arrived. Theo said hello.\n\n***\n\nA second scene begins. It ends here.\n\nChapter 2: Night\n\nFinal sentence.")
    extract(client, project)
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert len(chapters) == 2 and chapters[0]["status"] == "structured"
    scenes = client.get(f"/api/v1/chapters/{chapters[0]['id']}/scenes").json()
    assert len(scenes) == 2
    segments = client.get(f"/api/v1/scenes/{scenes[0]['id']}/segments").json()
    assert all(item["textContent"][-1] in ".!?" for item in segments)
    assert any(item["speakerCandidate"] == "Theo" for item in segments)


def test_extract_structure_after_failed_pdf_import_reports_canonical_source_requirement(client) -> None:
    project = client.post(
        "/api/v1/projects", json={"title": "Broken PDF", "rightsStatus": "declared"}
    ).json()["id"]
    import_job = client.post(
        f"/api/v1/projects/{project}/source/import",
        files={"file": ("broken.pdf", b"not a pdf", "application/pdf")},
        data={"rightsAcknowledged": "true"},
    ).json()
    assert wait_for_job(client, import_job["id"])["status"] == "failed"

    response = client.post(
        f"/api/v1/projects/{project}/structure/extract", json={"maxSegmentChars": 120}
    )
    assert response.status_code == 202
    job = wait_for_job(client, response.json()["id"])

    assert job["status"] == "failed"
    assert "A successfully imported canonical source is required" in job["errorMessage"]
    assert "NoneType" not in job["errorMessage"]
    assert "Capture a debug bundle" not in job["errorMessage"]


def test_unresolved_structure_and_segment_revision_history(client) -> None:
    project = project_with_source(client, "A single paragraph with no heading. Another complete sentence.")
    extract(client, project)
    chapter = client.get(f"/api/v1/projects/{project}/chapters").json()[0]
    assert chapter["status"] == "unresolved"
    scene = client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()[0]
    segment = client.get(f"/api/v1/scenes/{scene['id']}/segments").json()[0]
    edited = client.patch(f"/api/v1/segments/{segment['id']}", json={"textContent": "Corrected sentence."}).json()
    assert edited["revision"] == 2 and edited["status"] == "needs_review"
    revisions = client.get(f"/api/v1/segments/{segment['id']}/revisions").json()
    assert revisions[0]["revision"] == 1


def test_structure_parser_v2_front_matter_dialogue_and_warnings(client) -> None:
    project = project_with_source(
        client,
        "Dedication\n\n# Prologue\n\n[softly]\n\n\"Hello,\" she said.\n\nScene 2\n\nMara: We go.\n\nChapter 1: Start\n\nA final sentence.",
    )
    extract(client, project)

    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters][:2] == ["Dedication", "Prologue"]
    assert chapters[0]["status"] == "front_matter"
    warnings = client.get(f"/api/v1/projects/{project}/structure-warnings").json()
    assert any("Dialogue segment" in warning["message"] for warning in warnings)

    scenes = client.get(f"/api/v1/chapters/{chapters[1]['id']}/scenes").json()
    segments = [item for scene in scenes for item in client.get(f"/api/v1/scenes/{scene['id']}/segments").json()]
    assert any(item["segmentType"] == "performance_beat" for item in segments)
    assert any(item["segmentType"] == "dialogue" and item["speakerCandidate"] == "Mara" for item in segments)
    assert all("parserEvidence" in item for item in segments)


def test_structure_depth_language_and_front_back_matter_evidence(client) -> None:
    project = project_with_source(
        client,
        "Dedication\n\n"
        "Para la familia y los amigos, con la memoria de una casa antigua.\n\n"
        "Chapter 1\n\n"
        "El viento cruzó la plaza con una luz extraña. La niña miró la puerta.\n\n"
        "About the Author\n\n"
        "Con la vida entre libros, la autora recuerda una ciudad pequeña.",
    )
    extract(client, project)

    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters] == [
        "Dedication",
        "Chapter 1",
        "About the Author",
    ]
    assert [chapter["status"] for chapter in chapters] == [
        "front_matter",
        "structured",
        "back_matter",
    ]
    assert chapters[0]["parserEvidence"]["matterType"] == "front_matter"
    assert chapters[2]["parserEvidence"]["matterType"] == "back_matter"
    story_evidence = chapters[1]["parserEvidence"]
    assert story_evidence["language"] == "es"
    assert story_evidence["languageConfidence"] > 0
    assert story_evidence["languageEvidence"]["reason"] == "marker_heuristic"

    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["detectedLanguage"] == "es"
    assert quality["detectedLanguageConfidence"] > 0


def test_chapter_title_line_and_h3_h4_headings(client) -> None:
    project = project_with_source(
        client,
        "Chapter One\n"
        "The Door\n\n"
        "A first scene sentence.\n\n"
        "### Interior Beat\n\n"
        "This should stay inside chapter one.\n\n"
        "#### Camera Note\n\n"
        "This also stays inside chapter one.\n\n"
        "Chapter Two\n\n"
        "The ending.",
    )
    extract(client, project)

    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters] == [
        "Chapter One - The Door",
        "Chapter Two",
    ]


def test_explicit_and_possible_scene_breaks_report_quality(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        "The opening has enough body text to count before the separator.\n\n"
        "***\n\n"
        "A second scene has enough body text before time shifts arrive.\n\n"
        "Later that night, the windows rattled.\n\n"
        "The aftermath continued.",
    )
    extract(client, project)

    chapter = client.get(f"/api/v1/projects/{project}/chapters").json()[0]
    scenes = client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()
    assert len(scenes) == 3
    assert any(scene["confidence"] >= 0.9 for scene in scenes)
    assert "scene.possible_break_detected" in warning_codes(client, project)
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["sceneCount"] == 3
    assert quality["possibleSceneBreakCount"] == 1


def test_quote_tags_apostrophes_and_mixed_paragraphs(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        "\"I don't know,\" Mary-Jane said. I'm ready.\n\n"
        "\"Are you?\" asked Dr. Sen.\n\n"
        "[whispering]",
    )
    extract(client, project)

    segments = all_segments(client, project)
    assert [segment["segmentType"] for segment in segments] == [
        "dialogue",
        "narration",
        "dialogue",
        "performance_beat",
    ]
    assert segments[0]["speakerCandidate"] == "Mary-Jane"
    assert segments[1]["textContent"] == "I'm ready."
    assert segments[2]["speakerCandidate"] == "Dr. Sen"
    assert segments[0]["parserEvidence"]["productionType"] == "dialogue_with_tag"


def test_unclosed_quote_warns_without_dropping_text(client) -> None:
    source = 'Chapter 1\n\n"I don\'t know what happened.'
    project = project_with_source(client, source)
    extract(client, project)

    segments = all_segments(client, project)
    assert "".join(segment["textContent"] for segment in segments)
    assert "I don't know what happened." in " ".join(segment["textContent"] for segment in segments)
    warning = next(
        warning
        for warning in structure_warnings(client, project)
        if warning["evidence"].get("code") == "segment.quote_unclosed"
    )
    assert warning["evidence"]["reviewAction"] == "inspect_segment"
    assert warning["evidence"]["textPreview"] == '"I don\'t know what happened.'
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["quoteUnclosedCount"] == 1


def test_multiple_speaker_segment_is_split_or_marked_for_review(client) -> None:
    project = project_with_source(
        client,
        'Chapter 1\n\n"Come here," Priya said. "No," Rahul said.',
    )
    extract(client, project)

    dialogue = [
        segment for segment in all_segments(client, project)
        if segment["segmentType"] == "dialogue"
    ]
    clean_single_speaker_segments = [
        segment
        for segment in dialogue
        if segment["status"] == "ready" and segment["speakerCandidate"] in {"Priya", "Rahul"}
    ]
    assert len(clean_single_speaker_segments) != 1
    assert len(dialogue) == 2 or "segment.multiple_speakers" in warning_codes(client, project)


def test_mixed_narration_dialogue_splits_into_renderable_segments(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        'Rahul looked away. "I don\'t believe you," Priya whispered. '
        "The rain kept falling.",
    )
    extract(client, project)

    segments = all_segments(client, project)
    assert [segment["segmentType"] for segment in segments] == [
        "narration",
        "dialogue",
        "narration",
    ]
    assert segments[1]["speakerCandidate"] == "Priya"
    assert "segment.mixed_dialogue_and_narration" not in warning_codes(client, project)


def test_ambiguous_alternating_dialogue_warns_and_stays_reviewable(client) -> None:
    project = project_with_source(
        client,
        'Chapter 1\n\n"Where were you?"\n\n"Outside."\n\n"With whom?"\n\n"No one."',
    )
    extract(client, project)

    dialogue = [
        segment for segment in all_segments(client, project)
        if segment["segmentType"] == "dialogue"
    ]
    assert len(dialogue) == 4
    assert all(segment["status"] == "needs_review" for segment in dialogue)
    assert "speaker.ambiguous_two_person_exchange" in warning_codes(client, project)


def test_atom_offset_validation_direct_handles_apostrophe_dialogue() -> None:
    compiler = StructureCompiler("project", "source", "structure-parser-0.4.0")
    text = 'Rahul said, "I\'m here."'
    atoms = compiler.atoms_for_scene(text, 0)
    assert validate_atom_offsets(text, 0, atoms).valid


def test_atom_offset_validation_detects_non_whitespace_gap() -> None:
    atom = TextAtom("atom_1", "narration", "Alpha", 0, 5, None, 0.0, 1.0, {})

    validation = validate_atom_offsets("Alpha beta.", 0, [atom])

    assert validation.valid is False
    assert "uncovered_source_text" in validation.errors
    assert validation.uncovered_ranges == [(5, 11)]


def test_atom_offset_validation_allows_whitespace_only_gap() -> None:
    atoms = [
        TextAtom("atom_1", "narration", "Alpha", 0, 5, None, 0.0, 1.0, {}),
        TextAtom("atom_2", "narration", "beta", 8, 12, None, 0.0, 1.0, {}),
    ]

    validation = validate_atom_offsets("Alpha   beta", 0, atoms)

    assert validation.valid is True
    assert validation.uncovered_ranges == []


def test_atom_offset_validation_reports_overlap() -> None:
    atoms = [
        TextAtom("atom_1", "narration", "abc", 0, 3, None, 0.0, 1.0, {}),
        TextAtom("atom_2", "narration", "cde", 2, 5, None, 0.0, 1.0, {}),
        TextAtom("atom_3", "narration", "f", 5, 6, None, 0.0, 1.0, {}),
    ]

    validation = validate_atom_offsets("abcdef", 0, atoms)

    assert validation.valid is False
    assert "overlapping_atoms" in validation.errors
    assert validation.overlapping_ranges == [(2, 3)]


def test_atom_offset_validation_reports_slice_mismatch() -> None:
    atom = TextAtom("atom_1", "narration", "Alphi", 0, 5, None, 0.0, 1.0, {})

    validation = validate_atom_offsets("Alpha", 0, [atom])

    assert validation.valid is False
    assert "source_slice_mismatch" in validation.errors


def test_atom_offset_validation_reports_out_of_bounds() -> None:
    atom = TextAtom("atom_1", "narration", "Alpha!", 0, 6, None, 0.0, 1.0, {})

    validation = validate_atom_offsets("Alpha", 0, [atom])

    assert validation.valid is False
    assert "out_of_bounds" in validation.errors


def test_offset_validation_warning_includes_ranges(client, monkeypatch) -> None:
    original = StructureCompiler.atoms_for_scene

    def dropped_atom(self, scene_text, base, warnings=None, scene_id=None):
        atoms = original(self, scene_text, base, warnings, scene_id)
        return atoms[:1]

    monkeypatch.setattr(StructureCompiler, "atoms_for_scene", dropped_atom)
    project = project_with_source(client, "Chapter 1\n\nAlpha. Beta.")

    extract(client, project)

    warning = next(
        warning
        for warning in structure_warnings(client, project)
        if warning["evidence"].get("code") == "segment.offset_validation_failed"
    )
    evidence = warning["evidence"]
    assert evidence["reviewAction"] == "inspect_segment"
    assert "uncovered_source_text" in evidence["errors"]
    assert evidence["uncoveredRanges"]
    assert evidence["startOffset"] < evidence["endOffset"]


def test_alternating_unattributed_dialogue_remains_reviewable(client) -> None:
    project = project_with_source(client, "Chapter 1\n\n\"Yes.\"\n\n\"No.\"")
    extract(client, project)

    segments = all_segments(client, project)
    dialogue = [segment for segment in segments if segment["segmentType"] == "dialogue"]
    assert len(dialogue) == 2
    assert all(segment["status"] == "needs_review" for segment in dialogue)
    assert "segment.dialogue_no_speaker" in warning_codes(client, project)


def test_structure_depth_routes_multi_paragraph_dialogue_and_footnotes(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        "\"I began the story here,\n\n"
        "and I finished it after the pause,\" Mara said.\n\n"
        "[1] This source note should be inspected before narration.\n\n"
        "The room settled.",
    )
    extract(client, project)

    segments = all_segments(client, project)
    dialogue = [segment for segment in segments if segment["segmentType"] == "dialogue"]
    assert len(dialogue) == 1
    assert "and I finished it after the pause" in dialogue[0]["textContent"]
    assert dialogue[0]["status"] == "needs_review"
    assert "multi_paragraph_dialogue" in dialogue[0]["parserEvidence"]["atomReasons"]

    footnotes = [
        segment
        for segment in segments
        if segment["parserEvidence"]["productionType"] == "footnote"
    ]
    assert len(footnotes) == 1
    assert footnotes[0]["segmentType"] == "narration"
    assert footnotes[0]["status"] == "needs_review"
    assert footnotes[0]["parserEvidence"]["reviewAction"] == "inspect_footnote"
    assert "segment.footnote_routed" in footnotes[0]["parserEvidence"]["warningCodes"]
    assert "segment.footnote_routed" in warning_codes(client, project)
    assert "segment.offset_validation_failed" not in warning_codes(client, project)
    assert any(segment["textContent"] == "Mara said." for segment in segments)


def test_structure_depth_prosody_clause_splitting() -> None:
    compiler = StructureCompiler("project", "source", "structure-parser-0.4.0")
    source = (
        "Chapter 1\n\n"
        "The corridor held its breath, the lamp trembled against the wall, "
        "the old floor answered with a careful creak; the door remained closed, "
        "and Mara waited without speaking until the clock struck midnight."
    )

    result = compiler.compile(source, max_chars=80)

    segments = [
        segment
        for chapter in result.hierarchy
        for scene in chapter["scenes"]
        for segment in scene["segments"]
    ]
    narration = [segment for segment in segments if segment["segment_type"] == "narration"]
    assert len(narration) >= 3
    assert all(len(segment["text_content"]) <= 80 for segment in narration)
    assert any(segment["text_content"].endswith((",", ";")) for segment in narration[:-1])
    assert any(
        "prosody_clause_split"
        in json.loads(segment["parser_evidence_json"])["atomReasons"]
        for segment in narration
    )
    assert all(
        json.loads(str(warning["evidence_json"])).get("code")
        != "segment.offset_validation_failed"
        for warning in result.warnings
    )


def test_segment_split_merge_and_lock_survives_reextract(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\nA first sentence that is long enough to split near the middle. A second sentence follows for merging.",
    )
    extract(client, project)
    chapter = client.get(f"/api/v1/projects/{project}/chapters").json()[0]
    scene = client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()[0]
    segments = client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    target = segments[0]

    split_offset = target["textContent"].index(" enough")
    split = client.post(
        f"/api/v1/segments/{target['id']}/split", json={"splitOffset": split_offset}
    )
    assert split.status_code == 200
    split_segments = client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    assert len(split_segments) == len(segments) + 1

    merged = client.post(
        f"/api/v1/segments/{split_segments[0]['id']}/merge",
        json={"nextSegmentId": split_segments[1]["id"]},
    )
    assert merged.status_code == 200
    merged_segment = merged.json()
    locked = client.put(
        f"/api/v1/structure-locks/segment/{merged_segment['id']}",
        json={"locked": True, "reason": "Keep editorial split"},
    ).json()
    assert locked["userLocked"] is True

    extract(client, project)
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    scenes = [scene for chapter in chapters for scene in client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()]
    all_segments = [
        segment
        for scene in scenes
        for segment in client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    ]
    assert any(
        segment["id"] == merged_segment["id"] and segment["userLocked"]
        for segment in all_segments
    )


def test_llm_structure_refinement_creates_cast_and_speaker_rows(client, monkeypatch) -> None:
    monkeypatch.setattr(
        structure_module.StructureService,
        "_local_llm_ready",
        lambda _self: (True, "ready"),
    )
    monkeypatch.setattr(
        cast_discovery_module.CastDiscoveryService,
        "_local_llm_ready",
        lambda _self: False,
    )

    def fake_extract(_self, _project_id, request, _job_id=None, **_kwargs):
        if request.task == "atom_segment_refinement":
            atom_ids = atom_ids_from_prompt(request.prompt)
            return SimpleNamespace(
                run=SimpleNamespace(id="llmrun_segments"),
                result={
                    "segments": [
                        {
                            "atomIds": atom_ids[:2],
                            "segmentType": "dialogue",
                            "speakerHint": "Mara",
                            "confidence": 0.93,
                            "evidence": "speaker label",
                        },
                        {
                            "atomIds": atom_ids[2:],
                            "segmentType": "narration",
                            "speakerHint": "",
                            "confidence": 0.9,
                            "evidence": "narrative sentence",
                        },
                    ],
                    "warnings": [],
                },
            )
        return SimpleNamespace(run=SimpleNamespace(id="llmrun_empty"), result={"attributions": [], "warnings": []})

    monkeypatch.setattr(structure_module.LocalLlmService, "extract", fake_extract)
    project = project_with_source(
        client,
        'Chapter 1\n\n"We leave now," Mara said. The rain answered.',
    )

    extract(client, project)

    chapter = client.get(f"/api/v1/projects/{project}/chapters").json()[0]
    scene = client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()[0]
    segments = client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    assert [segment["segmentType"] for segment in segments] == ["dialogue", "narration"]
    assert segments[0]["speakerCandidate"] == "Mara"
    assert "optional_atom_llm_grouping" in segments[0]["parserEvidence"]["sources"]

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    mara = next(character for character in characters if character["displayName"] == "Mara")
    attributions = client.get(f"/api/v1/projects/{project}/speaker-attributions").json()
    mara_row = next(item for item in attributions if item["speakerName"] == "Mara")
    assert mara_row["status"] == "approved"
    assert mara_row["characterId"] == mara["id"]


def test_invalid_llm_structure_refinement_falls_back_with_warning(client, monkeypatch) -> None:
    monkeypatch.setattr(
        structure_module.StructureService,
        "_local_llm_ready",
        lambda _self: (True, "ready"),
    )
    monkeypatch.setattr(
        cast_discovery_module.CastDiscoveryService,
        "_local_llm_ready",
        lambda _self: True,
    )

    def fake_extract(_self, _project_id, request, _job_id=None, **_kwargs):
        if request.task == "atom_segment_refinement":
            atom_ids = atom_ids_from_prompt(request.prompt)
            return SimpleNamespace(
                run=SimpleNamespace(id="llmrun_bad_segments"),
                result={
                    "segments": [
                        {
                            "atomIds": atom_ids[:-1],
                            "segmentType": "dialogue",
                            "text": "Mara: We leave now. Invented text.",
                            "speakerHint": "Mara",
                            "confidence": 0.9,
                            "evidence": "bad output",
                        }
                    ],
                    "warnings": [],
                },
            )
        return SimpleNamespace(
            run=SimpleNamespace(id="llmrun_empty"),
            result={"characters": [], "decisions": [], "attributions": [], "warnings": []},
        )

    monkeypatch.setattr(structure_module.LocalLlmService, "extract", fake_extract)
    project = project_with_source(client, "Chapter 1\n\nMara: We leave now.")

    extract(client, project)

    chapter = client.get(f"/api/v1/projects/{project}/chapters").json()[0]
    scene = client.get(f"/api/v1/chapters/{chapter['id']}/scenes").json()[0]
    segments = client.get(f"/api/v1/scenes/{scene['id']}/segments").json()
    assert [segment["textContent"] for segment in segments] == ["Mara: We leave now."]
    warnings = client.get(f"/api/v1/projects/{project}/structure-warnings").json()
    assert any("failed validation" in warning["message"] for warning in warnings)


@pytest.mark.parametrize(
    "case",
    ["missing", "duplicate", "out_of_order", "invented", "non_adjacent"],
)
def test_invalid_atom_groupings_are_rejected(client, monkeypatch, case) -> None:
    monkeypatch.setattr(
        structure_module.StructureService,
        "_local_llm_ready",
        lambda _self: (True, "ready"),
    )
    monkeypatch.setattr(
        cast_discovery_module.CastDiscoveryService,
        "_local_llm_ready",
        lambda _self: False,
    )

    def fake_extract(_self, _project_id, request, _job_id=None, **_kwargs):
        if request.task == "atom_segment_refinement":
            atom_ids = atom_ids_from_prompt(request.prompt)
            if case == "missing":
                payload = [{"atomIds": atom_ids[:-1], "segmentType": "narration"}]
            elif case == "duplicate":
                payload = [
                    {"atomIds": atom_ids[:2], "segmentType": "dialogue", "speakerHint": "Mara"},
                    {"atomIds": atom_ids[1:], "segmentType": "narration"},
                ]
            elif case == "out_of_order":
                payload = [{"atomIds": [atom_ids[1], atom_ids[0]], "segmentType": "dialogue"}]
            elif case == "invented":
                payload = [{"atomIds": ["atom_not_real"], "segmentType": "narration"}]
            else:
                payload = [{"atomIds": [atom_ids[0], atom_ids[2]], "segmentType": "narration"}]
            return SimpleNamespace(
                run=SimpleNamespace(id=f"llmrun_bad_{case}"),
                result={"segments": payload, "warnings": []},
            )
        return SimpleNamespace(
            run=SimpleNamespace(id="llmrun_empty"),
            result={"characters": [], "decisions": [], "attributions": [], "warnings": []},
        )

    monkeypatch.setattr(structure_module.LocalLlmService, "extract", fake_extract)
    project = project_with_source(
        client,
        'Chapter 1\n\n"One," Mara said. The room held still. "Two," Theo said.',
    )

    extract(client, project)

    assert "llm.validation_failed" in warning_codes(client, project)
    assert all("Invented" not in segment["textContent"] for segment in all_segments(client, project))


def test_cast_candidates_and_speaker_rows_from_parser_evidence(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        "Dr. Sen: Begin.\n\n"
        "Mary-Jane: Wait.\n\n"
        "Captain Arjun: Hold.\n\n"
        "Mother: Come here.",
    )
    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    names = {character["displayName"] for character in characters}
    assert {"Dr. Sen", "Mary-Jane", "Captain Arjun", "Mother"} <= names
    attributions = client.get(f"/api/v1/projects/{project}/speaker-attributions").json()
    approved = {row["speakerName"]: row for row in attributions if row["status"] == "approved"}
    assert {"Dr. Sen", "Mary-Jane", "Captain Arjun", "Mother"} <= set(approved)


def test_possible_duplicate_cast_name_creates_review_issue(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nMary-Jane: Wait.")
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Mary"},
    ).json()

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    assert [character["displayName"] for character in characters] == [existing["displayName"]]
    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(issue for issue in issues if issue["title"] == "Possible duplicate cast candidate")
    assert issue["metadata"]["code"] == "cast.possible_duplicate"
    assert issue["metadata"]["reviewAction"] == "merge_cast"
    assert issue["metadata"]["candidateName"] == "Mary-Jane"
    assert issue["metadata"]["possibleMatches"] == ["Mary"]


def test_apply_merge_cast_issue_action_merges_candidate_and_resolves_issue(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nMary-Jane: Wait.")
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Mary"},
    ).json()

    extract(client, project)

    issue = next(
        issue
        for issue in client.get(f"/api/v1/projects/{project}/issues").json()
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    )
    response = client.post(
        f"/api/v1/issues/{issue['id']}/apply-action",
        json={"targetCharacterId": existing["id"]},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["result"]["action"] == "merge_cast"
    assert payload["result"]["characterId"] == existing["id"]
    assert payload["issue"]["status"] == "resolved"

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    target = next(character for character in characters if character["id"] == existing["id"])
    assert "Mary-Jane" in target["aliases"]
    assert any(
        character["displayName"] == "Mary-Jane"
        and character["mergedIntoCharacterId"] == existing["id"]
        for character in characters
    )


def test_speaker_proposal_uses_shortlist_review_for_honorific_alias(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\nDr. Priya Sen: Sit down.\n\nDr. Sen: Please listen.",
    )

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert [character["displayName"] for character in active] == ["Priya Sen"]
    assert "Dr. Priya Sen" in active[0]["aliases"]
    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(
        issue
        for issue in issues
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    )
    assert issue["metadata"]["candidateName"] == "Dr. Sen"
    assert issue["metadata"]["reviewAction"] == "merge_cast"
    assert issue["metadata"]["possibleMatches"] == ["Priya Sen"]
    assert issue["metadata"]["source"] == "speaker_attribution"
    assert issue["metadata"]["mentionCount"] == 1
    assert issue["metadata"]["windowIds"]
    assert issue["metadata"]["decisionId"]
    assert issue["metadata"]["evidenceGraph"]["canonicalName"] == "Dr. Sen"
    decisions = cast_graph_decisions(client, project)
    assert any(
        decision.source_name == "Dr. Sen"
        and decision.decision == "unsure"
        for decision in decisions
    )
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["possibleDuplicateCastCount"] == 1

    extract(client, project)

    issues_after_rerun = client.get(f"/api/v1/projects/{project}/issues").json()
    duplicate_issues = [
        issue
        for issue in issues_after_rerun
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    ]
    assert len(duplicate_issues) == 1

    client.patch(f"/api/v1/issues/{issue['id']}", json={"status": "resolved"})
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["possibleDuplicateCastCount"] == 0


def test_llm_mentions_additively_enrich_existing_character_observations(
    client, monkeypatch
) -> None:
    monkeypatch.setattr(
        cast_discovery_module.CastDiscoveryService, "_local_llm_ready", lambda _self: True
    )

    def fake_extract(_self, _project_id, request, _job_id=None, **_kwargs):
        if request.task != "cast_discovery":
            return SimpleNamespace(run=SimpleNamespace(id="llmrun_empty"), result={"warnings": []})
        segment_id = request.prompt.split("- ", 1)[1].split(" ", 1)[0]
        return SimpleNamespace(
            run=SimpleNamespace(id=None),
            result={
                "mentions": [
                    {
                        "surfaceName": "Priya",
                        "canonicalGuess": "Priya",
                        "entityType": "person",
                        "roleInScene": "speaker",
                        "evidenceText": "Priya: Speak carefully.",
                        "segmentIds": [segment_id],
                        "confidence": 0.78,
                        "traitsObserved": ["accent:irish"],
                        "relationshipsObserved": [
                            {"target": "Jon", "relation": "mentor", "confidence": 0.72}
                        ],
                        "speakingStyleObserved": ["measured"],
                    }
                ],
                "warnings": [],
            },
        )

    monkeypatch.setattr(structure_module.LocalLlmService, "extract", fake_extract)

    project = project_with_source(
        client,
        "Chapter 1\n\nPriya: Speak carefully.",
    )
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Priya", "confidence": 0.6},
    ).json()

    extract(client, project)
    cast_discovery_module.CastDiscoveryService(client.app.state.container).discover(
        project, use_local_llm=True
    )

    refreshed = client.get(f"/api/v1/projects/{project}/characters").json()
    priya = next(character for character in refreshed if character["id"] == existing["id"])
    assert "accent:irish" in priya["traits"]
    assert priya["confidence"] >= 0.78

    record = client.app.state.container.casting.character(existing["id"])
    assert record is not None
    assert json.loads(record.relationships_json) == [
        {"target": "Jon", "relation": "mentor", "confidence": 0.72}
    ]
    assert json.loads(record.speaking_style_json) == ["measured"]
    notes = json.loads(record.notes or "{}")
    assert notes["relationships"] == [{"target": "Jon", "relation": "mentor", "confidence": 0.72}]
    assert notes["speakingStyle"] == ["measured"]


def test_fuzzy_name_variant_routes_to_duplicate_review(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nElisabeth: Wait.")
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Elizabeth"},
    ).json()

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert [character["displayName"] for character in active] == [existing["displayName"]]

    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(
        issue
        for issue in issues
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    )
    assert issue["metadata"]["candidateName"] == "Elisabeth"
    assert issue["metadata"]["possibleMatches"] == ["Elizabeth"]


def test_duplicate_exact_display_names_require_explicit_review_target(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nAlex: Hold.")
    first = client.post(f"/api/v1/projects/{project}/characters", json={"displayName": "Alex"}).json()
    second = client.post(f"/api/v1/projects/{project}/characters", json={"displayName": "Alex"}).json()

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert [character["id"] for character in active] == [first["id"], second["id"]]

    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(
        issue
        for issue in issues
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    )
    assert issue["metadata"]["candidateName"] == "Alex"
    assert issue["metadata"]["possibleMatches"] == ["Alex", "Alex"]
    assert set(issue["metadata"]["possibleMatchIds"]) == {first["id"], second["id"]}
    assert issue["metadata"]["reason"]
    assert issue["metadata"]["decisionId"]


def test_generated_title_alias_enriches_existing_character_aliases(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nCaptain John: Stand down.")
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "John"},
    ).json()

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert [character["displayName"] for character in active] == ["John"]
    assert active[0]["id"] == existing["id"]
    assert "Captain John" in active[0]["aliases"]

    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    assert not [
        issue for issue in issues if issue["metadata"].get("code") == "cast.possible_duplicate"
    ]


def test_nickname_alias_clusters_without_duplicate(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nElizabeth: Wait.\n\nLiz: Go.")

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert len(active) == 1
    assert {"Elizabeth", "Liz"} <= {active[0]["displayName"], *active[0]["aliases"]}

    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    duplicate_issues = [
        issue
        for issue in issues
        if issue["metadata"].get("code") == "cast.possible_duplicate"
    ]
    assert duplicate_issues == []


def test_nickname_siblings_cluster_without_canonical_mention(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nLiz: Go.\n\nBeth: Wait.")

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert len(active) == 1
    assert {"Beth", "Elizabeth", "Liz"} <= {active[0]["displayName"], *active[0]["aliases"]}


def test_transitive_nickname_aliases_refresh_discovery_index(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\nRob: Go.\n\nRobert: Wait.\n\nBob: Listen.",
    )

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert len(active) == 1
    assert {"Rob", "Robert", "Bob"} <= {active[0]["displayName"], *active[0]["aliases"]}


def test_cast_discovery_extracts_title_aliases_and_traits(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\n"
        "Captain Mara: Hold the line.\n\n"
        "The young Irish captain raised her hand.\n\n"
        "Mara: We move at dawn.",
    )

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [character for character in characters if not character["mergedIntoCharacterId"]]
    assert [character["displayName"] for character in active] == ["Mara"]
    character = active[0]
    assert "Captain Mara" in character["aliases"]
    assert "role:captain" in character["traits"]
    notes = json.loads(character["notes"])
    graph = notes["evidenceGraph"]
    assert graph["mentionEvidenceCount"] == 2
    assert graph["firstSeenOffset"] is not None
    assert graph["lastSeenOffset"] is not None


def test_cast_graph_filters_noise_mentions_and_writes_manifest(client, monkeypatch) -> None:
    monkeypatch.setattr(
        cast_discovery_module.CastDiscoveryService, "_local_llm_ready", lambda _self: True
    )

    def fake_extract(_self, _project_id, request, _job_id=None, **_kwargs):
        if request.task != "cast_discovery":
            return SimpleNamespace(run=SimpleNamespace(id="llmrun_empty"), result={"warnings": []})
        segment_ids = [
            line.split("- ", 1)[1].split(" ", 1)[0]
            for line in request.prompt.splitlines()
            if line.startswith("- ")
        ]
        return SimpleNamespace(
            run=SimpleNamespace(id=None),
            result={
                "mentions": [
                    {
                        "surfaceName": "PROJECT HAIL MARY",
                        "canonicalGuess": "PROJECT HAIL MARY",
                        "entityType": "title",
                        "roleInScene": "mentioned",
                        "evidenceText": "PROJECT HAIL MARY",
                        "segmentIds": [segment_ids[0]],
                        "confidence": 0.92,
                        "traitsObserved": [],
                        "relationshipsObserved": [],
                        "speakingStyleObserved": [],
                    },
                    {
                        "surfaceName": "Andy Weir",
                        "canonicalGuess": "Andy Weir",
                        "entityType": "author",
                        "roleInScene": "mentioned",
                        "evidenceText": "Andy Weir",
                        "segmentIds": [segment_ids[0]],
                        "confidence": 0.91,
                        "traitsObserved": [],
                        "relationshipsObserved": [],
                        "speakingStyleObserved": [],
                    },
                    {
                        "surfaceName": "Mara",
                        "canonicalGuess": "Mara",
                        "entityType": "person",
                        "roleInScene": "speaker",
                        "evidenceText": "Mara: Go.",
                        "segmentIds": [segment_ids[-1]],
                        "confidence": 0.94,
                        "traitsObserved": [],
                        "relationshipsObserved": [],
                        "speakingStyleObserved": [],
                    },
                ],
                "warnings": [],
            },
        )

    monkeypatch.setattr(structure_module.LocalLlmService, "extract", fake_extract)

    project = project_with_source(
        client,
        "Dedication\n\nPROJECT HAIL MARY\n\nANDY WEIR\n\nChapter 1\n\nMara: Go.",
    )

    extract(client, project)
    cast_discovery_module.CastDiscoveryService(client.app.state.container).discover(
        project, use_local_llm=True
    )

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active_names = {
        character["displayName"]
        for character in characters
        if not character["mergedIntoCharacterId"]
    }
    assert active_names == {"Mara"}

    mentions = cast_graph_mentions(client, project)
    by_surface = {
        mention.surface_name: [
            json.loads(item.metadata_json)
            for item in mentions
            if item.surface_name == mention.surface_name
        ]
        for mention in mentions
    }
    assert all(item["filteredOut"] is True for item in by_surface["PROJECT HAIL MARY"])
    assert all(item["filteredOut"] is True for item in by_surface["Andy Weir"])
    assert any(item["filteredOut"] is False for item in by_surface["Mara"])

    manifest = casting_manifest(client, project)
    assert manifest["payload"]["windowCount"] >= 1
    assert manifest["payload"]["mentionCount"] >= 3
    assert manifest["payload"]["filteredMentionCount"] >= 2
    assert manifest["payload"]["candidateCount"] >= 1
    assert manifest["payload"]["decisionCount"] >= 1
    assert any(
        item["surfaceName"] == "PROJECT HAIL MARY" and item["filteredOut"] is True
        for item in manifest["payload"]["mentions"]
    )


def test_low_confidence_cast_candidate_issue_metadata(client) -> None:
    project = project_with_source(client, 'Chapter 1\n\nRahul looked away. "Hello."')

    extract(client, project)

    assert "cast.low_confidence_candidate" in issue_codes(client, project)
    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(
        issue
        for issue in issues
        if issue["metadata"].get("code") == "cast.low_confidence_candidate"
    )
    assert issue["metadata"]["candidateName"] == "Rahul"
    assert issue["metadata"]["reviewAction"] == "confirm_cast"
    assert issue["metadata"]["confidence"] < 0.72
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["lowConfidenceCastCandidateCount"] == 1

    client.patch(f"/api/v1/issues/{issue['id']}", json={"status": "resolved"})
    quality = client.get(f"/api/v1/projects/{project}/structure/quality").json()
    assert quality["lowConfidenceCastCandidateCount"] == 0


def test_apply_confirm_cast_issue_action_creates_character_and_resolves_issue(client) -> None:
    project = project_with_source(client, 'Chapter 1\n\nRahul looked away. "Hello."')

    extract(client, project)

    issue = next(
        issue
        for issue in client.get(f"/api/v1/projects/{project}/issues").json()
        if issue["metadata"].get("code") == "cast.low_confidence_candidate"
    )
    response = client.post(f"/api/v1/issues/{issue['id']}/apply-action", json={})

    assert response.status_code == 200
    payload = response.json()
    assert payload["result"]["action"] == "confirm_cast"
    assert payload["issue"]["status"] == "resolved"
    assert payload["result"]["characterId"]

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    created = next(character for character in characters if character["displayName"] == "Rahul")
    assert created["id"] == payload["result"]["characterId"]
    assert created["confidence"] == issue["metadata"]["confidence"]


def test_apply_issue_action_rejects_missing_review_action(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nA clean sentence.")
    created = client.post(
        f"/api/v1/projects/{project}/issues",
        json={
            "category": "manual_review",
            "severity": "warning",
            "title": "Manual issue",
            "description": "No automation is attached.",
        },
    ).json()

    response = client.post(f"/api/v1/issues/{created['id']}/apply-action", json={})

    assert response.status_code == 422
    assert "reviewAction" in response.json()["detail"]


def test_cast_evidence_graph_counts_mentions(client) -> None:
    project = project_with_source(
        client,
        "Chapter 1\n\nPriya: Sit down.\n\nRahul looked at Priya.",
    )

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    priya = next(character for character in characters if character["displayName"] == "Priya")
    notes = json.loads(priya["notes"])
    graph = notes["evidenceGraph"]
    assert graph["speakerEvidenceCount"] >= 1
    assert graph["mentionEvidenceCount"] >= 1
    assert graph["firstSeenOffset"] is not None
    assert graph["lastSeenOffset"] is not None
    assert "mention" in graph["sources"]
    assert notes["mentionEvidence"]


def _docx_heading_bytes() -> bytes:
    document = Document()
    document.add_heading("The Arrival", level=1)
    document.add_paragraph("Mara stepped off the train into the cold morning air.")
    document.add_paragraph("She had never seen the city before that day.")
    document.add_heading("The Departure", level=1)
    document.add_paragraph("The whistle blew and the crowded platform slowly emptied.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _epub_heading_bytes(tmp_path: Path) -> bytes:
    book = epub.EpubBook()
    book.set_identifier("container-signals")
    book.set_title("Container Signals")
    book.set_language("en")
    first = epub.EpubHtml(title="The Arrival", file_name="c1.xhtml", lang="en", uid="arrival")
    first.content = (
        "<h1>The Arrival</h1><p>Mara stepped off the train into the cold morning air.</p>"
        "<p>She had never seen the city before that day.</p>"
    )
    second = epub.EpubHtml(title="The Departure", file_name="c2.xhtml", lang="en", uid="departure")
    second.content = (
        "<h1>The Departure</h1><p>The whistle blew and the crowded platform slowly emptied.</p>"
    )
    book.add_item(first)
    book.add_item(second)
    book.toc = (first, second)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", first, second]
    target = tmp_path / "container.epub"
    epub.write_epub(str(target), book)
    return target.read_bytes()


def _import_document(client, name: str, data: bytes) -> str:
    project = client.post(
        "/api/v1/projects", json={"title": "Container", "rightsStatus": "declared"}
    ).json()["id"]
    job = client.post(
        f"/api/v1/projects/{project}/source/import",
        files={"file": (name, data, "application/octet-stream")},
        data={"rightsAcknowledged": "true"},
    ).json()
    assert wait_for_job(client, job["id"])["status"] == "succeeded"
    return project


def test_docx_heading_styles_become_chapters(client) -> None:
    project = _import_document(client, "book.docx", _docx_heading_bytes())
    extract(client, project)
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters] == ["The Arrival", "The Departure"]
    assert all(chapter["status"] == "structured" for chapter in chapters)


def test_epub_spine_and_toc_become_chapters(client, tmp_path: Path) -> None:
    project = _import_document(client, "book.epub", _epub_heading_bytes(tmp_path))
    extract(client, project)
    chapters = client.get(f"/api/v1/projects/{project}/chapters").json()
    assert [chapter["title"] for chapter in chapters] == ["The Arrival", "The Departure"]


def test_container_signal_matches_heading_longer_than_display_title() -> None:
    heading = "An Uncommonly Long Heading About " + "Very " * 25 + "Distant Shores"
    assert len(heading) > 120
    text = f"{heading}\n\nA body paragraph follows the long heading here.\n"
    compiler = StructureCompiler("proj_test", "src_test", "test-parser")
    signal = ChapterSignal(
        title=heading[:120],
        source_kind="docx_heading",
        level=1,
        anchor_text=heading,
        confidence=0.95,
    )

    result = compiler.compile(text, 200, chapter_signals=[signal])

    titles = [chapter["record"]["title"] for chapter in result.hierarchy]
    assert titles == [heading[:120]]
    codes = {
        json.loads(str(warning["evidence_json"])).get("code") for warning in result.warnings
    }
    assert "container_signal_unmatched" not in codes


def test_cast_discovery_uses_aliases_without_creating_duplicates(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nCaptain Vale: Report.")
    character = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Mara", "aliases": ["Captain Vale"]},
    ).json()

    extract(client, project)

    characters = client.get(f"/api/v1/projects/{project}/characters").json()
    active = [item for item in characters if not item["mergedIntoCharacterId"]]
    assert [item["displayName"] for item in active] == ["Mara"]
    attributions = client.get(f"/api/v1/projects/{project}/speaker-attributions").json()
    row = next(item for item in attributions if item["speakerName"] == "Captain Vale")
    assert row["status"] == "approved"
    assert row["characterId"] == character["id"]


def _reparse(client, project: str) -> None:
    job = client.post(f"/api/v1/projects/{project}/source/reparse", json={}).json()
    assert wait_for_job(client, job["id"])["status"] == "succeeded"


def test_reject_merge_suppresses_duplicate_reflagging(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nMary-Jane: Wait.")
    existing = client.post(
        f"/api/v1/projects/{project}/characters",
        json={"displayName": "Mary"},
    ).json()
    extract(client, project)

    issues = client.get(f"/api/v1/projects/{project}/issues").json()
    issue = next(
        issue for issue in issues if issue["metadata"].get("code") == "cast.possible_duplicate"
    )
    assert issue["metadata"]["candidateName"] == "Mary-Jane"

    rejected = client.post(
        f"/api/v1/characters/{existing['id']}/reject-merge",
        json={"candidateName": "Mary-Jane", "reason": "Different people."},
    )
    assert rejected.status_code == 200

    resolved = client.get(f"/api/v1/projects/{project}/issues").json()
    resolved_issue = next(item for item in resolved if item["id"] == issue["id"])
    assert resolved_issue["status"] == "resolved"

    # Re-parse mints fresh segment ids (new dedupe keys); without the rejected
    # decision this would re-flag the same pair as a brand-new open issue.
    _reparse(client, project)
    extract(client, project)

    after = client.get(f"/api/v1/projects/{project}/issues").json()
    open_duplicates = [
        item
        for item in after
        if item["metadata"].get("code") == "cast.possible_duplicate"
        and item["metadata"].get("candidateName") == "Mary-Jane"
        and item["status"] == "open"
    ]
    assert open_duplicates == []


def test_reject_merge_validates_character_exists(client) -> None:
    missing = client.post(
        "/api/v1/characters/char_does_not_exist/reject-merge",
        json={"candidateName": "Mary-Jane"},
    )
    assert missing.status_code == 404


def test_merge_prompt_uses_recent_decisions_and_caps_shortlist(client) -> None:
    project = project_with_source(client, "Chapter 1\n\nQuiet narration only.")
    survivor = client.post(
        f"/api/v1/projects/{project}/characters", json={"displayName": "Bran Hale"}
    ).json()
    source = client.post(
        f"/api/v1/projects/{project}/characters", json={"displayName": "Bran Hall"}
    ).json()
    client.post(
        f"/api/v1/characters/{survivor['id']}/merge",
        json={"sourceCharacterId": source["id"], "reason": "Same person."},
    )
    for name in [
        "Brian Hail",
        "Brin Hail",
        "Brand Hail",
        "Brandon Hail",
        "Bran Hale",
        "Bran Hall",
    ]:
        client.post(f"/api/v1/projects/{project}/characters", json={"displayName": name})

    service = cast_discovery_module.CastDiscoveryService(client.app.state.container)
    candidate = cast_discovery_module.CharacterCandidate(
        display_name="Bran Hail",
        canonical_name="Bran Hail",
        aliases=[],
        first_seen_segment_id=None,
        first_seen_chapter_id=None,
        evidence=["{}"],
        role_guess="supporting",
        confidence=0.74,
        source="llm_cast_discovery",
        mention_evidence=["{}"],
    )

    shortlist = service._character_index(project).shortlist(candidate)
    prompt = service._merge_prompt(project, candidate, shortlist)

    assert len(shortlist) == 5
    assert "confirmed same person" in prompt
    assert prompt.count("id=") == 5


def test_structure_v2_chunking_and_coverage_helpers() -> None:
    text = "Chapter 1\n\n" + ("A paragraph with text.\n\n" * 20)
    chunks = structure_v2.chunk_text(text, chunk_chars=120, overlap_chars=20)

    assert len(chunks) > 1
    assert chunks[0].start_offset == 0
    assert chunks[1].read_only_context_start <= chunks[1].start_offset
    assert structure_v2.seam_windows(chunks)

    valid = structure_v2.verify_structure_coverage(
        [
            {
                "scenes": [
                    {
                        "segments": [
                            {"record": {"start_offset": 0, "end_offset": 5, "text_content": "Hello"}},
                            {"record": {"start_offset": 5, "end_offset": 11, "text_content": " world"}},
                        ]
                    }
                ]
            }
        ]
    )
    overlap = structure_v2.verify_structure_coverage(
        [
            {
                "scenes": [
                    {
                        "segments": [
                            {"record": {"start_offset": 0, "end_offset": 6, "text_content": "Hello!"}},
                            {"record": {"start_offset": 5, "end_offset": 11, "text_content": " world"}},
                        ]
                    }
                ]
            }
        ]
    )

    assert valid.ok is True
    assert overlap.ok is False
    assert overlap.overlap_count == 1


def test_structure_v2_flag_adds_manifest_trace(client, settings) -> None:
    flagged_settings = settings.__class__(
        **{**settings.__dict__, "structure_v2_enabled": True}
    )
    with TestClient(create_app(flagged_settings)) as flagged_client:
        project = project_with_source(
            flagged_client,
            "Chapter 1\n\nMara opened the door.\n\nChapter 2\n\nMara closed it.",
        )
        extract(flagged_client, project)
        project_record = flagged_client.app.state.container.projects.get(project)
        assert project_record
        manifest_path = Path(project_record.artifact_path) / "manifests" / "structure_manifest.json"
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

    trace = manifest["payload"]["structureV2"]
    assert trace["version"] == structure_v2.STRUCTURE_V2_VERSION
    assert trace["coverage"]["segmentCount"] > 0
    assert trace["fallback"] == "deterministic_compiler"
